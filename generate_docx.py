from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- LETTERHEAD ---
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = header.add_run("VIDYA SETU INTERNATIONAL\n")
title_run.bold = True
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(31, 38, 135)

subtitle_run = header.add_run("123 Innovation Tech Park, Bengaluru, Karnataka, 560001\ncontact@vidyasetu.com | +91-800-123-4567\n")
subtitle_run.font.size = Pt(10)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# --- TITLE ---
title = doc.add_heading('Vidya Setu Architecture & Business Compliance Report', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# --- IMPORTANCE ---
doc.add_heading('1. The Importance of Vidya Setu', level=2)
doc.add_paragraph(
    "Vidya Setu is a paradigm shift in Indian educational administration. Unlike legacy ERPs (e.g. Fedena, Entab), "
    "which were built as rigid, bloated accounting systems that treated data security as an afterthought, Vidya Setu is "
    "designed fundamentally around the user. It prioritizes fluid communication, strict data protection, and robust compliance."
)
doc.add_paragraph(
    "By moving away from monolithic data structures and separating modules based on strict role-based access control (RBAC), "
    "Vidya Setu completely prevents data bleeding. Parents see only their children; teachers see only their classes; and highly "
    "sensitive information (like POCSO grievance reports and medical history) is completely inaccessible to unauthorized staff."
)

# --- ARCHITECTURE & RULES ---
doc.add_heading('2. System Architecture & Compliance Rules', level=2)
doc.add_paragraph("The platform is engineered around the following unshakeable rules:")

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Digital Personal Data Protection (DPDP) Act Compliance: ").bold = True
p1.add_run("All user data is explicitly tracked with Consent Records. If a parent requests data deletion (Erasure Request), the system correctly soft-deletes the identity while legally preserving necessary financial audit logs.")

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("POCSO Act Compliant Grievance Handling: ").bold = True
p2.add_run("Child safety is paramount. The database includes a dedicated Grievance Box where sensitive reports are directly forwarded to the designated Child Protection Committee Member, bypassing standard teacher/clerk visibility completely.")

p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Decoupled Identity Profiles: ").bold = True
p3.add_run("A 'User' is strictly an authentication identity. Their actual data lives in isolated profile tables (StudentProfile, TeacherProfile). A teacher cannot accidentally gain parent-level access, and a student's medical file (StudentHealthRecord) is strictly separated from their academic marks.")

# --- ISO STANDARDS ---
doc.add_heading('3. ISO Standards Implementation', level=2)
doc.add_paragraph(
    "Vidya Setu's backend and database architecture is built specifically to meet the auditing requirements of global standards:"
)

iso1 = doc.add_paragraph(style='List Bullet')
iso1.add_run("ISO/IEC 27001 (Information Security Management): ").bold = True
iso1.add_run("Implemented through forced JWT authentication, invisible reCAPTCHA for bot protection, and an immutable AuditLog table that tracks every single user login, data modification, and failed authentication attempt.")

iso2 = doc.add_paragraph(style='List Bullet')
iso2.add_run("ISO/IEC 27701 (Privacy Information Management): ").bold = True
iso2.add_run("Implemented via Data Subject Request (DSR) tracking. Users can legally request their data, and the platform retains a cryptographically verifiable log of their consent.")

iso3 = doc.add_paragraph(style='List Bullet')
iso3.add_run("ISO 9001 (Quality Management Systems): ").bold = True
iso3.add_run("The DocumentApproval workflow ensures that critical academic changes and certificates require a digital cryptographic signature from the Principal before issuance, maintaining an unbroken chain of custody.")

# --- UNIT ECONOMICS ---
doc.add_heading('4. Clean, Transparent Unit Economics', level=2)
doc.add_paragraph(
    "Vidya Setu operates on a highly scalable B2B2C Software-as-a-Service (SaaS) model. Our cloud-native architecture "
    "(Next.js + AWS EC2 + PostgreSQL) allows us to maintain incredibly low server overhead while maximizing profit margins."
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value (INR)'
hdr_cells[2].text = 'Details'

data = [
    ("Revenue Per Student (ARPU)", "₹250 - ₹500 / year", "Billed to the school annually per enrolled student."),
    ("Cloud Infrastructure Cost", "₹15 - ₹20 / student", "AWS Hosting, PostgreSQL RDS, SMS Gateway (OTP)."),
    ("Customer Acquisition Cost (CAC)", "₹15,000 / school", "Sales team overhead to onboard a school (average 1,000 students)."),
    ("Gross Margin", "85% - 92%", "Highly profitable due to automated multi-tenant architecture."),
    ("Payback Period", "< 2 Months", "The ₹2,50,000+ annual contract instantly covers the CAC.")
]

for item, val, desc in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = val
    row_cells[2].text = desc

doc.add_paragraph()
doc.add_paragraph("Because the platform is fully multi-tenant, onboarding a new school requires zero engineering effort. We simply provision a new School ID in the database and they are instantly live. The unit economics prove that Vidya Setu is a highly lucrative, rapidly scaling enterprise.")

doc.save('/Users/sumanthemmanuel/Desktop/VidyaSetu_Business_Report.docx')
